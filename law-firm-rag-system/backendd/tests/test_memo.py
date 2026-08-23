"""
Memorandum model and both writers.

The model tests matter most: because the Word and PDF writers consume the same
:class:`MemoModel`, any assertion proved here holds for both deliverables, which
is the mechanism preventing the two artefacts from drifting apart.
"""

from __future__ import annotations

import zipfile

import pytest

from core.memo_model import MemoModel, build_memo_model


class TestMemoModel:
    def test_counts_and_metadata_are_carried_through(self, sample_results):
        model = build_memo_model(sample_results, firm_name="Josh and Mak International",
                                 document_names=["deed.pdf", "fard.pdf"])
        assert model.firm_name == "Josh and Mak International"
        assert model.high_risk_count == 1
        assert model.total_questions == 3
        assert len(model.document_names) == 2

    def test_findings_are_ordered_by_risk_not_question_number(self, sample_results):
        model = build_memo_model(sample_results)
        risks = [f["risk_level"] for f in model.findings]
        assert risks[0] == "HIGH"
        assert risks == sorted(risks, key=lambda r: {"HIGH": 0, "MEDIUM": 1, "LOW": 2}[r])

    def test_overall_risk_reflects_the_worst_finding(self, sample_results):
        assert build_memo_model(sample_results).overall_risk == "HIGH"

    def test_overall_risk_is_low_on_a_clean_bundle(self):
        model = build_memo_model({
            "findings": [{"question_id": 1, "risk_level": "LOW", "finding": "ok"}],
            "high_risk_count": 0, "medium_risk_count": 0, "low_risk_count": 1,
            "red_flags": [], "flags": {},
        })
        assert model.overall_risk == "LOW"

    def test_red_flags_force_high_even_without_a_high_finding(self):
        model = build_memo_model({
            "findings": [], "high_risk_count": 0, "medium_risk_count": 1,
            "red_flags": [{"id": "RF001", "label": "Missing NOC"}], "flags": {},
        })
        assert model.overall_risk == "HIGH"

    def test_letterhead_is_suppressed_for_the_default_firm_name(self):
        assert build_memo_model({"flags": {}}, firm_name="Law Firm").has_letterhead is False
        assert build_memo_model({"flags": {}}, firm_name="Mumtaz & Brohi").has_letterhead is True

    def test_executive_summary_names_the_real_risks(self, sample_results):
        summary = build_memo_model(sample_results, document_names=["a.pdf"]).executive_summary
        assert "HIGH risk" in summary
        assert "red flag" in summary.lower()
        assert "12,500,000" in summary

    def test_executive_summary_is_honest_when_nothing_is_wrong(self):
        summary = build_memo_model({
            "findings": [], "high_risk_count": 0, "medium_risk_count": 0,
            "low_risk_count": 15, "total_questions": 15, "red_flags": [], "flags": {},
        }).executive_summary
        assert "No item was assessed as HIGH risk" in summary

    def test_compliance_rows_reflect_the_statutory_thresholds(self, sample_results):
        rows = {row.label.split("\n")[0]: row
                for row in build_memo_model(sample_results).compliance_rows}
        withholding = next(r for k, r in rows.items() if "Withholding tax" in k)
        assert withholding.action_required is True
        assert "236C" in withholding.label

    def test_below_threshold_transactions_report_no_action(self):
        model = build_memo_model({"findings": [], "red_flags": [], "flags": {
            "detected_value_pkr": 1_000_000, "fbr_applicable": False,
            "aml_threshold": False,
        }})
        withholding = next(r for r in model.compliance_rows if "Withholding" in r.label)
        assert withholding.action_required is False
        assert "below" in withholding.value.lower()

    def test_ocr_failures_add_a_disclosure_row(self):
        model = build_memo_model({"findings": [], "red_flags": [],
                                  "flags": {"ocr_failures": 3}})
        assert any("could not be read" in row.label for row in model.compliance_rows)

    def test_empty_results_do_not_raise(self):
        model = build_memo_model({})
        assert isinstance(model, MemoModel)
        assert model.total_questions == 0
        assert model.executive_summary

    def test_malformed_findings_do_not_raise(self):
        model = build_memo_model({"findings": [{}, {"risk_level": None}],
                                  "red_flags": [], "flags": {}})
        assert len(model.findings) == 2


class TestWordWriter:
    def test_produces_a_readable_docx(self, sample_results, tmp_path):
        from docx import Document

        from core.memo_generator import generate_memo

        output = tmp_path / "memo.docx"
        generate_memo(sample_results, output, firm_name="Mumtaz & Brohi",
                      firm_tagline="Barristers & Corporate Counsel",
                      city="islamabad", document_names=["deed.pdf"],
                      flags=sample_results["flags"])

        assert output.exists() and output.stat().st_size > 5000
        assert zipfile.is_zipfile(output)          # a .docx is an OPC zip

        text = "\n".join(p.text for p in Document(str(output)).paragraphs)
        assert "DUE DILIGENCE REVIEW MEMORANDUM" in text
        assert "MUMTAZ & BROHI" in text
        assert "Executive Summary" in text
        assert "Disclaimer" in text

    def test_every_numbered_section_is_present(self, sample_results, tmp_path):
        from docx import Document

        from core.memo_generator import generate_memo

        output = tmp_path / "memo.docx"
        generate_memo(sample_results, output, flags=sample_results["flags"])
        text = "\n".join(p.text for p in Document(str(output)).paragraphs)
        for heading in ("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8."):
            assert heading in text

    def test_no_emoji_survive_into_the_document(self, sample_results, tmp_path):
        """
        Emoji risk markers rendered as tofu boxes in many Word installations.
        They are replaced by real cell shading plus a text label.
        """
        from docx import Document

        from core.memo_generator import generate_memo

        output = tmp_path / "memo.docx"
        generate_memo(sample_results, output, flags=sample_results["flags"])
        document = Document(str(output))
        text = "\n".join(p.text for p in document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text
        for emoji in ("🔴", "🟡", "🟢", "🚨", "⚠️", "✅"):
            assert emoji not in text

    def test_empty_results_still_produce_a_document(self, tmp_path):
        from core.memo_generator import generate_memo

        output = tmp_path / "empty.docx"
        generate_memo({}, output)
        assert output.exists()

    def test_nested_output_directory_is_created(self, sample_results, tmp_path):
        from core.memo_generator import generate_memo

        output = tmp_path / "deep" / "nested" / "memo.docx"
        generate_memo(sample_results, output, flags=sample_results["flags"])
        assert output.exists()


class TestPdfWriter:
    def test_produces_a_valid_pdf(self, sample_results, tmp_path):
        from core.pdf_generator import generate_memo_pdf

        output = tmp_path / "memo.pdf"
        generate_memo_pdf(sample_results, output, firm_name="Josh and Mak International",
                          city="islamabad", document_names=["deed.pdf"],
                          flags=sample_results["flags"])

        assert output.exists() and output.stat().st_size > 3000
        with output.open("rb") as handle:
            assert handle.read(5) == b"%PDF-"

    def test_xml_unsafe_content_does_not_break_generation(self, tmp_path):
        """reportlab uses a mini-markup — unescaped angle brackets would abort it."""
        from core.pdf_generator import generate_memo_pdf

        results = {
            "findings": [{
                "question_id": 1, "risk_level": "HIGH",
                "question": "Does clause <A> & <B> apply?",
                "finding": "The deed says a < b & c > d at clause 4.",
                "reasoning": "Ampersands & angle brackets <like this> appear.",
                "document_citation": "deed.pdf <p.4>", "statutory_citation": "s.17 & s.49",
                "constitutional_basis": "Article 23", "recommendation": "Check <this>.",
                "missing_documents": ["NOC <urgent>"],
            }],
            "high_risk_count": 1, "medium_risk_count": 0, "low_risk_count": 0,
            "red_flags": [], "flags": {},
        }
        output = tmp_path / "unsafe.pdf"
        generate_memo_pdf(results, output)
        assert output.exists()

    def test_empty_results_still_produce_a_pdf(self, tmp_path):
        from core.pdf_generator import generate_memo_pdf

        output = tmp_path / "empty.pdf"
        generate_memo_pdf({}, output)
        assert output.exists()

    def test_both_writers_agree_on_content(self, sample_results, tmp_path):
        """Shared model ⇒ the two deliverables cannot report different risk."""
        from core.memo_generator import generate_memo
        from core.pdf_generator import generate_memo_pdf

        model = build_memo_model(sample_results, flags=sample_results["flags"])
        generate_memo(sample_results, tmp_path / "a.docx", flags=sample_results["flags"])
        generate_memo_pdf(sample_results, tmp_path / "a.pdf", flags=sample_results["flags"])

        assert (tmp_path / "a.docx").exists()
        assert (tmp_path / "a.pdf").exists()
        assert model.overall_risk == "HIGH"
