"""
Word (.docx) memorandum writer.

Improvements over the previous build:

* **Emoji were used as risk markers** (🔴 🟡 🟢).  Word substitutes a fallback
  glyph for these in many installed fonts, so a memorandum that looked correct
  on the developer's machine printed as tofu boxes on a client's.  Risk is now
  conveyed by real cell shading plus a text label, which renders identically
  everywhere and survives monochrome printing.
* **Cells were only ever font-coloured**, never shaded, because ``python-docx``
  exposes no shading API — the previous build simply did without.  Shading is
  applied here through the underlying ``w:shd`` element.
* **Findings appeared in question order**, so a HIGH-risk item could sit on
  page nine.  They are now ordered by risk, then by question number.
* **Page numbers, an executive summary and a document inventory were absent.**
  All three are now present, and every value is defensively coerced so a
  malformed finding cannot abort generation of the whole memorandum.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from core.memo_model import DISCLAIMER, MemoModel, build_memo_model

logger = logging.getLogger(__name__)

# ── Palette ───────────────────────────────────────────────────────────────
NAVY = RGBColor(0x14, 0x3C, 0x6E)
AZURE = RGBColor(0x00, 0x70, 0xA4)
RISK_RED = RGBColor(0xAA, 0x23, 0x23)
AMBER = RGBColor(0xB8, 0x5C, 0x00)
FOREST = RGBColor(0x1E, 0x6E, 0x37)
SLATE = RGBColor(0x4B, 0x55, 0x63)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

SHADE_HEADER = "143C6E"
SHADE_LABEL = "EEF2F8"
SHADE_HIGH = "FBE9E9"
SHADE_MEDIUM = "FDF2E3"
SHADE_LOW = "EAF6EE"
SHADE_MUTED = "F4F6FA"

RISK_TEXT = {
    "HIGH": ("HIGH RISK", RISK_RED, SHADE_HIGH),
    "MEDIUM": ("MEDIUM RISK", AMBER, SHADE_MEDIUM),
    "LOW": ("LOW RISK", FOREST, SHADE_LOW),
}


# ──────────────────────────────────────────────────────────────────────────
#  Low-level helpers
# ──────────────────────────────────────────────────────────────────────────
def _shade(cell: Any, hex_colour: str) -> None:
    """Apply background shading — python-docx has no API for this."""
    try:
        element = OxmlElement("w:shd")
        element.set(qn("w:val"), "clear")
        element.set(qn("w:color"), "auto")
        element.set(qn("w:fill"), hex_colour)
        cell._tc.get_or_add_tcPr().append(element)
    except Exception:                                        # noqa: BLE001
        pass


def _write_cell(cell: Any, text: str, *, bold: bool = False,
                colour: RGBColor | None = None, size: int = 9,
                shade: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(str(text if text is not None else ""))
    run.bold = bold
    run.font.size = Pt(size)
    if colour is not None:
        run.font.color.rgb = colour
    if shade:
        _shade(cell, shade)


def _heading(doc: Document, text: str, *, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = NAVY


def _rule(doc: Document, colour: RGBColor = NAVY) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run("─" * 78)
    run.font.size = Pt(6)
    run.font.color.rgb = colour


def _body(doc: Document, text: str, *, size: int = 10,
          colour: RGBColor | None = None, italic: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    if colour is not None:
        run.font.color.rgb = colour


def _add_page_numbers(section: Any) -> None:
    """Insert a PAGE / NUMPAGES field pair into the footer."""
    try:
        paragraph = section.footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def _field(instruction: str) -> None:
            begin = OxmlElement("w:fldChar")
            begin.set(qn("w:fldCharType"), "begin")
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = instruction
            end = OxmlElement("w:fldChar")
            end.set(qn("w:fldCharType"), "end")
            run = paragraph.add_run()
            run._r.append(begin)
            run._r.append(instr)
            run._r.append(end)
            run.font.size = Pt(8)
            run.font.color.rgb = SLATE

        prefix = paragraph.add_run("Page ")
        prefix.font.size = Pt(8)
        prefix.font.color.rgb = SLATE
        _field("PAGE")
        middle = paragraph.add_run(" of ")
        middle.font.size = Pt(8)
        middle.font.color.rgb = SLATE
        _field("NUMPAGES")
    except Exception:                                        # noqa: BLE001
        logger.debug("Could not add page-number field (non-fatal).")


def _table(doc: Document, rows: int, cols: int) -> Any:
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    return table


# ──────────────────────────────────────────────────────────────────────────
#  Sections
# ──────────────────────────────────────────────────────────────────────────
def _letterhead(doc: Document, model: MemoModel) -> None:
    if not model.has_letterhead:
        return
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(2)
    run = name.add_run(model.firm_name.upper())
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = NAVY

    if model.firm_tagline:
        tagline = doc.add_paragraph()
        tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tagline.paragraph_format.space_after = Pt(2)
        run = tagline.add_run(model.firm_tagline)
        run.italic = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = AZURE

    contact = " · ".join(
        part for part in (model.firm_address, model.firm_phone, model.firm_email) if part
    )
    if contact:
        line = doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line.paragraph_format.space_after = Pt(4)
        run = line.add_run(contact)
        run.font.size = Pt(8.5)
        run.font.color.rgb = SLATE
    _rule(doc)


def _cover(doc: Document, model: MemoModel) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run(model.title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run(model.subtitle)
    run.font.size = Pt(12)
    run.font.color.rgb = AZURE

    label, colour, shade = RISK_TEXT.get(model.overall_risk, RISK_TEXT["LOW"])
    banner = _table(doc, 1, 1)
    _write_cell(banner.rows[0].cells[0],
                f"OVERALL ASSESSMENT: {label}",
                bold=True, colour=colour, size=11, shade=shade)
    banner.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _summary_section(doc: Document, model: MemoModel) -> None:
    _heading(doc, "1.  Transaction Summary")
    rows = [
        ("Prepared by", model.firm_name),
        ("Date of review", model.generated_at),
        ("Transaction type", model.transaction_type.title()),
        ("City", model.city.title()),
        ("Controlling authority",
         model.authority_full_name or "Not determined"),
        ("Applicable bye-laws", model.relevant_bylaws or "Not determined"),
        ("Housing society", model.housing_society or "Not applicable"),
        ("Documents reviewed", str(len(model.document_names))),
        ("Questions assessed", str(model.total_questions)),
    ]
    table = _table(doc, len(rows), 2)
    for index, (label, value) in enumerate(rows):
        _write_cell(table.rows[index].cells[0], label, bold=True, colour=NAVY, shade=SHADE_LABEL)
        _write_cell(table.rows[index].cells[1], value)
    doc.add_paragraph()

    if model.document_names:
        _heading(doc, "1.1  Documents in the reviewed bundle", level=2)
        for name in model.document_names:
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(name)
            run.font.size = Pt(9.5)
        doc.add_paragraph()


def _executive_summary(doc: Document, model: MemoModel) -> None:
    _heading(doc, "2.  Executive Summary")
    _body(doc, model.executive_summary)
    doc.add_paragraph()


def _risk_section(doc: Document, model: MemoModel) -> None:
    _heading(doc, "3.  Risk Summary")
    table = _table(doc, 2, 4)
    headers = ["HIGH RISK", "MEDIUM RISK", "LOW RISK", "NOT ASSESSED"]
    shades = [SHADE_HIGH, SHADE_MEDIUM, SHADE_LOW, SHADE_MUTED]
    colours = [RISK_RED, AMBER, FOREST, SLATE]
    counts = [model.high_risk_count, model.medium_risk_count,
              model.low_risk_count, model.failed_count]

    for column, (header, shade, colour, count) in enumerate(
            zip(headers, shades, colours, counts)):
        head_cell = table.rows[0].cells[column]
        _write_cell(head_cell, header, bold=True, colour=colour, size=9, shade=shade)
        head_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_cell = table.rows[1].cells[column]
        _write_cell(value_cell, f"{count} item{'s' if count != 1 else ''}",
                    bold=True, colour=colour, size=14, shade=shade)
        value_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _red_flags_section(doc: Document, model: MemoModel) -> None:
    _heading(doc, "4.  Red Flags")
    if not model.red_flags:
        _body(doc, "No transaction-stopping defect was detected in the documents "
                   "supplied. This is not a warranty of clean title; it records only "
                   "that the deterministic rules in this system did not fire.",
              colour=FOREST)
        doc.add_paragraph()
        return

    table = _table(doc, len(model.red_flags) + 1, 3)
    for column, header in enumerate(["Reference", "Defect", "Statutory and constitutional basis"]):
        _write_cell(table.rows[0].cells[column], header,
                    bold=True, colour=WHITE, shade=SHADE_HEADER)
    for index, flag in enumerate(model.red_flags, start=1):
        row = table.rows[index]
        _write_cell(row.cells[0], str(flag.get("id", "")), bold=True,
                    colour=RISK_RED, shade=SHADE_HIGH)
        detail = str(flag.get("label", ""))
        if flag.get("description"):
            detail += f"\n{flag['description']}"
        _write_cell(row.cells[1], detail, colour=RISK_RED)
        basis = f"{flag.get('statute', '')}\n{flag.get('article', '')}".strip()
        _write_cell(row.cells[2], basis, size=8, colour=SLATE)
    doc.add_paragraph()


def _findings_section(doc: Document, model: MemoModel) -> None:
    _heading(doc, "5.  Clause-by-Clause Findings")
    _body(doc, "Findings are ordered by assessed risk. Each entry records the evidence "
               "relied on and the provision it engages, so that any conclusion can be "
               "independently verified against the original document.",
          size=9, colour=SLATE, italic=True)

    if not model.findings:
        _body(doc, "No findings were produced.", colour=SLATE)
        return

    for finding in model.findings:
        risk = str(finding.get("risk_level", "LOW")).upper()
        label, colour, shade = RISK_TEXT.get(risk, RISK_TEXT["LOW"])

        header = doc.add_paragraph()
        header.paragraph_format.space_before = Pt(10)
        header.paragraph_format.space_after = Pt(2)
        header.paragraph_format.keep_with_next = True
        run = header.add_run(f"Q{finding.get('question_id', '—')}.  "
                             f"{str(finding.get('question', ''))[:170]}")
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = NAVY

        rows = [
            ("Risk", label),
            ("Finding", finding.get("finding")),
            ("Reasoning", finding.get("reasoning")),
            ("Document citation", finding.get("document_citation")),
            ("Statutory citation", finding.get("statutory_citation")),
            ("Constitutional basis", finding.get("constitutional_basis")),
            ("Recommendation", finding.get("recommendation")),
        ]
        missing = finding.get("missing_documents") or []
        if missing:
            rows.append(("Documents required", "; ".join(str(m) for m in missing)))
        confidence = finding.get("confidence")
        if confidence:
            rows.append(("Model confidence", str(confidence).upper()))

        table = _table(doc, len(rows), 2)
        for index, (key, value) in enumerate(rows):
            _write_cell(table.rows[index].cells[0], key,
                        bold=True, colour=NAVY, size=8.5, shade=SHADE_LABEL)
            if key == "Risk":
                _write_cell(table.rows[index].cells[1], str(value),
                            bold=True, colour=colour, shade=shade)
            else:
                _write_cell(table.rows[index].cells[1],
                            str(value if value else "Not stated"))
    doc.add_paragraph()


def _compliance_section(doc: Document, model: MemoModel) -> None:
    _heading(doc, "6.  Tax, AML and Processing Compliance")
    table = _table(doc, len(model.compliance_rows) + 1, 2)
    _write_cell(table.rows[0].cells[0], "Matter", bold=True, colour=WHITE, shade=SHADE_HEADER)
    _write_cell(table.rows[0].cells[1], "Determination", bold=True, colour=WHITE,
                shade=SHADE_HEADER)

    for index, row in enumerate(model.compliance_rows, start=1):
        _write_cell(table.rows[index].cells[0], row.label,
                    bold=True, colour=NAVY, size=8.5, shade=SHADE_LABEL)
        text = row.value + (f"\n{row.note}" if row.note else "")
        _write_cell(
            table.rows[index].cells[1], text,
            colour=RISK_RED if row.action_required else None,
            bold=row.action_required,
            shade=SHADE_HIGH if row.action_required else None,
        )
    doc.add_paragraph()
    _body(doc,
          "Statutory reference: Income Tax Ordinance 2001, sections 236C and 236K "
          "(withholding tax on immovable property transactions); Anti-Money Laundering "
          "Act 2010 (enhanced due diligence); SECP AML/CFT Regulations 2018 (designated "
          "non-financial businesses and professions). Constitutional basis: Article 23 — "
          "right to acquire property subject to law.",
          size=8.5, colour=SLATE, italic=True)
    doc.add_paragraph()


def _missing_section(doc: Document, model: MemoModel) -> None:
    _heading(doc, "7.  Documents to Requisition")
    if not model.missing_documents:
        _body(doc, "No further documents were identified as missing from the bundle.",
              colour=FOREST)
        doc.add_paragraph()
        return
    for name in model.missing_documents:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(str(name))
        run.font.size = Pt(9.5)
        run.font.color.rgb = RISK_RED
    doc.add_paragraph()


def _disclaimer_section(doc: Document) -> None:
    _heading(doc, "8.  Scope and Disclaimer")
    _body(doc, DISCLAIMER, size=9, colour=SLATE)


# ──────────────────────────────────────────────────────────────────────────
#  Entry point
# ──────────────────────────────────────────────────────────────────────────
def generate_memo(
    results: dict[str, Any],
    output_path: str | Path,
    firm_name: str = "Law Firm",
    firm_address: str = "",
    firm_phone: str = "",
    firm_email: str = "",
    firm_tagline: str = "",
    transaction_type: str = "property",
    city: str = "islamabad",
    document_names: list[str] | None = None,
    flags: dict[str, Any] | None = None,
) -> str:
    """Render the memorandum as a Word document and return its path."""
    model = build_memo_model(
        results,
        firm_name=firm_name, firm_address=firm_address, firm_phone=firm_phone,
        firm_email=firm_email, firm_tagline=firm_tagline,
        transaction_type=transaction_type, city=city,
        document_names=document_names, flags=flags,
    )
    return write_docx(model, output_path)


def write_docx(model: MemoModel, output_path: str | Path) -> str:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.2)
        section.start_type = WD_SECTION.NEW_PAGE

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    header_paragraph = doc.sections[0].header.paragraphs[0]
    header_paragraph.text = f"{model.firm_name} — Due Diligence Review"
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if header_paragraph.runs:
        header_paragraph.runs[0].font.size = Pt(8)
        header_paragraph.runs[0].font.color.rgb = SLATE
    _add_page_numbers(doc.sections[0])

    _letterhead(doc, model)
    _cover(doc, model)
    _summary_section(doc, model)
    _executive_summary(doc, model)
    _risk_section(doc, model)
    _red_flags_section(doc, model)
    _findings_section(doc, model)
    _compliance_section(doc, model)
    _missing_section(doc, model)
    _disclaimer_section(doc)

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    logger.info("Word memorandum written: %s", path)
    return str(path)
