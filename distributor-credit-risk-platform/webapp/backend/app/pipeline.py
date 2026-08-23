"""
Shared Scoring Pipeline Module
Problem 01: Distributor Credit Risk on Gut Feel

This module contains the SAME logic as score_dealers.py and robust_ingestion.py,
refactored into importable functions operating on in-memory DataFrames instead
of fixed file paths -- necessary because the web app accepts uploaded files
rather than reading from disk.

CRITICAL: every function here was verified to reproduce EXACT known values
from the validated CLI pipeline (D0080 = 418, 220 dealers, 178/36/6 split)
before this module was considered correct. See verify_against_cli.py.
"""

import io
import os
import re
import math
from datetime import date, timedelta
from functools import lru_cache
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# The model was trained on features computed from history ending at this date.
# Kept as the default so the original dataset scores identically -- but it is
# only meaningful FOR that dataset, which is why resolve_cutoff_date() below
# falls back to a data-derived cutoff for any other distributor's upload.
DEFAULT_CUTOFF_DATE = pd.Timestamp(os.environ.get("FEATURE_CUTOFF_DATE", "2025-01-01"))


def resolve_cutoff_date(txns: pd.DataFrame) -> tuple[pd.Timestamp, dict]:
    """
    Decides which date separates 'payment history used for features' from
    everything after it.

    A fixed calendar date breaks on any data that sits entirely after it --
    the feature window comes back empty and the caller sees a confusing
    "no dealers could be scored" error with no indication why.

    Rule:
      - If the default falls strictly INSIDE the uploaded data's date range,
        use it. (For the original dataset this provably routes to the same
        code path as before, so existing scores are unchanged.)
      - Otherwise derive it from the data: one day after the latest
        transaction, so the entire uploaded history counts as features.

    Returns (cutoff_date, info) -- info is surfaced in the API response so
    the caller can see which rule fired instead of guessing.
    """
    valid = txns["due_date"].dropna()
    if valid.empty:
        return DEFAULT_CUTOFF_DATE, {
            "cutoff_date": str(DEFAULT_CUTOFF_DATE.date()),
            "strategy": "default (no valid dates found in upload)",
        }

    data_min, data_max = valid.min(), valid.max()

    if data_min < DEFAULT_CUTOFF_DATE < data_max:
        return DEFAULT_CUTOFF_DATE, {
            "cutoff_date": str(DEFAULT_CUTOFF_DATE.date()),
            "strategy": "configured default (falls within uploaded data range)",
            "data_range": f"{data_min.date()} to {data_max.date()}",
        }

    derived = data_max + pd.Timedelta(days=1)
    return derived, {
        "cutoff_date": str(derived.date()),
        "strategy": "derived from uploaded data (default cutoff falls outside data range)",
        "data_range": f"{data_min.date()} to {data_max.date()}",
    }
# Currency erosion used for inflation-adjusted exposure. The 0.18 default
# reflects recent PKR conditions -- it is wrong outside Pakistan and drifts
# wrong within it as inflation changes, so it is configurable.
ANNUAL_PKR_EROSION = float(os.environ.get("ANNUAL_CURRENCY_EROSION", 0.18))
MIN_INVOICES = 3

REASON_LABELS = {
    "payment_delay_severity": "Late and inconsistent payment timing",
    "bounce_rate_lifetime": "Cheque bounce history",
    "real_exposure_pkr": "Inflation-adjusted credit exposure",
    "order_frequency_trend": "Declining order activity",
    "salesman_default_rate_loo": "Salesman's track record with similar dealers",
    "territory_default_rate_loo": "Elevated risk in dealer's territory",
}

TRUE_SET = {"y", "yes", "1", "true", "bounced"}
FALSE_SET = {"n", "no", "0", "false", "", "ok", "nan"}

# --- Eid / Ramzan seasonal windows -----------------------------------------
# The original hardcoded table only covered Mar 2023 - Jun 2025. Any newer
# data silently got NOTHING flagged as seasonal -- no error, the feature just
# stopped working. These explicit windows are preserved verbatim so the
# original dataset scores identically; every OTHER year is computed from the
# Islamic calendar at runtime.
#
# Uses the tabular (arithmetic) Islamic calendar -- no external dependency,
# so there is nothing extra to install or fail at deploy time. Validated
# against 9 real observed Ramadan/Eid dates (2023-2025): max deviation 1 day,
# which is immaterial against +/- 2 week windows. Note that observed dates
# depend on local moon sighting and Pakistan often differs from Saudi Arabia
# by a day; the window padding absorbs that.

_EXPLICIT_SEASONAL_WINDOWS = [
    (date(2023, 3, 10), date(2023, 5, 5)), (date(2023, 6, 15), date(2023, 7, 10)),
    (date(2024, 2, 28), date(2024, 4, 25)), (date(2024, 6, 5), date(2024, 6, 30)),
    (date(2025, 2, 15), date(2025, 4, 15)), (date(2025, 5, 25), date(2025, 6, 20)),
]
_EXPLICIT_YEARS = {2023, 2024, 2025}

RAMADAN_LEAD_DAYS = 14      # window opens before Ramadan (stock-up period)
EID_FITR_TAIL_DAYS = 15     # window closes after Eid-ul-Fitr (collection lag)
EID_AZHA_PAD_DAYS = 13      # symmetric padding around Eid-ul-Azha


def _gregorian_to_jdn(y, m, d):
    a = (14 - m) // 12
    y2 = y + 4800 - a
    m2 = m + 12 * a - 3
    return d + (153 * m2 + 2) // 5 + 365 * y2 + y2 // 4 - y2 // 100 + y2 // 400 - 32045


def _jdn_to_gregorian(jdn):
    a = jdn + 32044
    b = (4 * a + 3) // 146097
    c = a - (146097 * b) // 4
    d2 = (4 * c + 3) // 1461
    e = c - (1461 * d2) // 4
    m2 = (5 * e + 2) // 153
    day = e - (153 * m2 + 2) // 5 + 1
    month = m2 + 3 - 12 * (m2 // 10)
    year = 100 * b + d2 - 4800 + m2 // 10
    return date(year, month, day)


def _hijri_to_jdn(y, m, d):
    return (d + math.ceil(29.5 * (m - 1)) + (y - 1) * 354
            + math.floor((3 + 11 * y) / 30) + 1948440 - 1)


def _hijri_to_date(y, m, d):
    return _jdn_to_gregorian(_hijri_to_jdn(y, m, d))


def _hijri_year_of(g: date) -> int:
    jdn = _gregorian_to_jdn(g.year, g.month, g.day)
    return int(math.floor((30 * (jdn - 1948440) + 10646) / 10631))


@lru_cache(maxsize=256)
def _computed_windows_for_hijri_year(h_year):
    """Ramadan/Eid-ul-Fitr window and Eid-ul-Azha window for one Hijri year."""
    ramadan_start = _hijri_to_date(h_year, 9, 1)    # 1 Ramadan
    eid_fitr      = _hijri_to_date(h_year, 10, 1)   # 1 Shawwal
    eid_azha      = _hijri_to_date(h_year, 12, 10)  # 10 Dhu al-Hijjah
    return (
        (ramadan_start - timedelta(days=RAMADAN_LEAD_DAYS),
         eid_fitr + timedelta(days=EID_FITR_TAIL_DAYS)),
        (eid_azha - timedelta(days=EID_AZHA_PAD_DAYS),
         eid_azha + timedelta(days=EID_AZHA_PAD_DAYS)),
    )


def is_seasonal(d):
    if pd.isna(d):
        return False
    d = d.date() if hasattr(d, "date") else d

    # Explicit table wins for the years it covers -- guarantees the original
    # dataset produces byte-identical seasonal flags.
    for start, end in _EXPLICIT_SEASONAL_WINDOWS:
        if start <= d <= end:
            return True
    if d.year in _EXPLICIT_YEARS:
        return False

    # Ramadan drifts ~11 days earlier each Gregorian year, so a window can
    # straddle a year boundary -- check the neighbouring Hijri years too.
    hy = _hijri_year_of(d)
    for y in (hy - 1, hy, hy + 1):
        for start, end in _computed_windows_for_hijri_year(y):
            if start <= d <= end:
                return True
    return False


# ---------------------------------------------------------------------------
# INGESTION (ported from robust_ingestion.py)
# ---------------------------------------------------------------------------
def parse_messy_date(val):
    if pd.isna(val) or str(val).strip() in ("", "N/A", "NA", "-", "nil", "pending"):
        return pd.NaT
    s = str(val).strip()
    if s.isdigit() and 40000 < int(s) < 48000:
        try:
            return pd.Timestamp("1899-12-30") + pd.Timedelta(days=int(s))
        except Exception:
            return pd.NaT
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%m-%d-%Y", "%d-%b-%y", "%Y.%m.%d", "%d-%b-%Y"):
        try:
            return pd.to_datetime(s, format=fmt)
        except Exception:
            continue
    # ISO-like strings are unambiguous; only apply dayfirst to genuinely
    # ambiguous formats, which is what raises the pandas warning otherwise.
    try:
        return pd.to_datetime(s, format="ISO8601", errors="raise")
    except Exception:
        pass
    try:
        return pd.to_datetime(s, dayfirst=True, errors="raise")
    except Exception:
        return pd.NaT


def parse_messy_amount(val):
    if pd.isna(val):
        return np.nan
    s = str(val).strip()
    if s in ("", "N/A", "NA", "-", "nil"):
        return np.nan
    s = s.strip("()").replace("Rs.", "").replace("Rs", "").replace("PKR", "")
    s = s.replace(",", "").replace("/-", "").strip()
    try:
        return float(s)
    except ValueError:
        return np.nan


def parse_bool(val):
    s = str(val).strip().lower()
    if s in TRUE_SET:
        return True
    if s in FALSE_SET:
        return False
    return np.nan


def clean_transactions(raw_bytes: bytes, valid_dealer_ids: set) -> tuple[pd.DataFrame, dict]:
    """
    Cleans a raw (potentially messy) uploaded transactions CSV.
    Returns (cleaned_dataframe, quality_report_dict).
    Same rules as robust_ingestion.py -- NaN detection uses .isna() explicitly
    (not string-sentinel matching), per the bug fix caught during the
    Day 2 stress test.
    """
    raw = pd.read_csv(io.BytesIO(raw_bytes), dtype=str)
    raw.columns = [c.strip() for c in raw.columns]

    df, _ = apply_column_aliases(raw, TRANSACTION_COLUMN_ALIASES, "transactions")
    df = df[[c for c in df.columns if not c.startswith("Unnamed") and c != "Notes"]]

    report = {"input_rows": len(df)}

    df = df.drop_duplicates()

    missing_dealer_raw = df["dealer_id"].isna() if "dealer_id" in df.columns else pd.Series([False] * len(df))
    df["dealer_id"] = df["dealer_id"].astype(str).str.strip()
    missing_dealer = missing_dealer_raw | df["dealer_id"].isin(["nan", "", "None", "N/A", "NA"])
    n_missing = int(missing_dealer.sum())
    df = df[~missing_dealer]

    orphan_mask = ~df["dealer_id"].isin(valid_dealer_ids)
    n_orphan = int(orphan_mask.sum())
    df = df[~orphan_mask]

    for col in ["invoice_date", "due_date", "payment_date"]:
        df[col] = df[col].apply(parse_messy_date)
    before = len(df)
    df = df[df["invoice_date"].notna() & df["due_date"].notna()]
    n_bad_dates = before - len(df)

    df["amount_pkr"] = df["amount_pkr"].apply(parse_messy_amount)
    before = len(df)
    df = df[df["amount_pkr"].notna() & (df["amount_pkr"] > 0)]
    n_bad_amount = before - len(df)

    df["cheque_bounced"] = df["cheque_bounced"].apply(parse_bool)
    before = len(df)
    df = df[df["cheque_bounced"].notna()]
    n_bad_bool = before - len(df)

    before = len(df)
    df = df[df["due_date"] >= df["invoice_date"]]
    n_bad_logic = before - len(df)

    df["days_late"] = (df["payment_date"] - df["due_date"]).dt.days
    df.loc[df["cheque_bounced"] == True, "days_late"] = np.nan
    df["is_eid_ramzan_period"] = df["due_date"].apply(is_seasonal)

    final_cols = ["transaction_id", "dealer_id", "invoice_date", "due_date", "payment_date",
                  "amount_pkr", "payment_method", "cheque_bounced", "days_late", "is_eid_ramzan_period"]
    for c in final_cols:
        if c not in df.columns:
            df[c] = np.nan
    clean = df[final_cols].reset_index(drop=True)

    report.update({
        "output_rows": len(clean),
        "retention_rate": round(len(clean) / report["input_rows"], 3) if report["input_rows"] else 0,
        "dropped_missing_dealer": n_missing,
        "dropped_orphan_dealer": n_orphan,
        "dropped_bad_dates": n_bad_dates,
        "dropped_bad_amount": n_bad_amount,
        "dropped_bad_bounce_flag": n_bad_bool,
        "dropped_logical_inconsistency": n_bad_logic,
    })
    return clean, report


# ---------------------------------------------------------------------------
# FEATURE ENGINEERING + SCORING (ported from score_dealers.py)
# ---------------------------------------------------------------------------
def _safe_zscore(s: pd.Series) -> pd.Series:
    """Z-score that degrades to 0 when a series has no spread.

    A single-dealer portfolio, or one where every dealer behaves identically,
    has zero (or undefined) variance. Dividing by it yields NaN, which travels
    silently through StandardScaler and then makes predict_proba raise
    'Input X contains NaN' -- surfacing as an unexplained 500. Zero is the
    correct answer here: with no spread, nobody deviates from the average.
    """
    std = s.std()
    if not np.isfinite(std) or std == 0:
        return pd.Series(0.0, index=s.index)
    return (s - s.mean()) / std

def loo_group_rate(df, group_col, rate_col):
    s = df.groupby(group_col)[rate_col].transform("sum")
    c = df.groupby(group_col)[rate_col].transform("count")
    return ((s - df[rate_col]) / (c - 1).replace(0, np.nan)).fillna(df[rate_col].mean())


def compute_features(dealers: pd.DataFrame, salesmen: pd.DataFrame, txns: pd.DataFrame,
                     cutoff_date: pd.Timestamp | None = None):
    """Returns (feat_df, insufficient_history_dealer_ids).

    cutoff_date is optional -- when omitted it is resolved from the data, so
    existing callers (regression tests, verify_against_cli.py) keep working
    unchanged."""
    if cutoff_date is None:
        cutoff_date, _ = resolve_cutoff_date(txns)
    dealers = dealers.copy()
    # Dealer dates get the same tolerant parsing the transactions file already
    # receives. A bare pd.to_datetime() infers a format from the first row and
    # then crashes on the first value that contradicts it -- which is exactly
    # what a DD/MM/YYYY export does, and DD/MM/YYYY is the norm in Pakistan.
    dealers["onboarding_date"] = dealers["onboarding_date"].apply(parse_messy_date)
    if dealers["onboarding_date"].isna().any():
        bad = dealers.loc[dealers["onboarding_date"].isna(), "dealer_id"].head(5).tolist()
        raise ValueError(
            "Could not read the onboarding date for some dealers "
            f"(e.g. {', '.join(map(str, bad))}). Accepted formats include "
            "YYYY-MM-DD, DD/MM/YYYY, MM-DD-YYYY and Excel serial numbers."
        )
    feature_window = txns[txns["due_date"] < cutoff_date]

    feat_rows, insufficient = [], []
    for dealer_id, dealer_row in dealers.set_index("dealer_id").iterrows():
        dgrp = feature_window[feature_window["dealer_id"] == dealer_id]
        if len(dgrp) < MIN_INVOICES:
            insufficient.append(dealer_id)
            continue
        non_bounced = dgrp[dgrp["cheque_bounced"] == False]
        bounce_rate_lifetime = dgrp["cheque_bounced"].mean()
        avg_days_late = non_bounced["days_late"].mean() if len(non_bounced) else 0
        nonseasonal = non_bounced[non_bounced.get("is_eid_ramzan_period", False) == False]
        avg_days_late_nonseasonal = nonseasonal["days_late"].mean() if len(nonseasonal) else avg_days_late
        payment_volatility = non_bounced["days_late"].std() if len(non_bounced) > 1 else 0
        years_active = max((cutoff_date.date() - dealer_row["onboarding_date"].date()).days / 365.25, 0)
        real_exposure_pkr = dealer_row["credit_limit_pkr"] * ((1 - ANNUAL_PKR_EROSION) ** years_active)
        mid = feature_window["invoice_date"].median()
        early = dgrp[dgrp["invoice_date"] < mid]
        late = dgrp[dgrp["invoice_date"] >= mid]
        order_frequency_trend = (len(late) - len(early)) / max(len(early), 1)
        feat_rows.append({
            "dealer_id": dealer_id, "dealer_name": dealer_row["dealer_name"],
            "city": dealer_row["city"], "sector": dealer_row["sector"],
            "salesman_id": dealer_row["salesman_id"], "territory_risk_tier": dealer_row["territory_risk_tier"],
            "is_salesman_favorite": bool(dealer_row["is_salesman_favorite"]),
            "credit_limit_pkr": dealer_row["credit_limit_pkr"],
            "bounce_rate_lifetime": bounce_rate_lifetime,
            "avg_days_late_nonseasonal": avg_days_late_nonseasonal,
            "payment_volatility": payment_volatility,
            "real_exposure_pkr": real_exposure_pkr,
            "order_frequency_trend": order_frequency_trend,
        })

    feat_df = pd.DataFrame(feat_rows)
    if len(feat_df) == 0:
        return feat_df, insufficient

    feat_df["salesman_default_rate_loo"] = loo_group_rate(feat_df, "salesman_id", "bounce_rate_lifetime")
    feat_df["territory_default_rate_loo"] = loo_group_rate(feat_df, "territory_risk_tier", "bounce_rate_lifetime")
    feat_df["payment_delay_severity"] = (
        _safe_zscore(feat_df["avg_days_late_nonseasonal"])
        + _safe_zscore(feat_df["payment_volatility"])
    ) / 2
    feat_df = feat_df.merge(salesmen[["salesman_id", "salesman_name"]], on="salesman_id", how="left")

    # Final safety net. Any residual NaN/inf -- e.g. a dealer whose payments
    # all have a missing payment_date -- would otherwise reach predict_proba
    # and raise an opaque "Input X contains NaN" 500. Zero is the neutral
    # value: on a z-scored feature it means "at the portfolio average", on a
    # rate feature it means "no observed events".
    _model_cols = ["payment_delay_severity", "bounce_rate_lifetime",
                   "real_exposure_pkr", "order_frequency_trend",
                   "salesman_default_rate_loo", "territory_default_rate_loo"]
    _present = [c for c in _model_cols if c in feat_df.columns]
    feat_df[_present] = (feat_df[_present]
                         .replace([np.inf, -np.inf], np.nan)
                         .fillna(0.0))

    return feat_df, insufficient


def score_with_model(feat_df: pd.DataFrame, artifact: dict) -> pd.DataFrame:
    """Applies the persisted model/scaler exactly as score_dealers.py does. No .fit() calls."""
    model, scaler, FEATURE_COLS = artifact["model"], artifact["scaler"], artifact["feature_columns"]
    score_params = artifact["score_params"]

    X_score_s = scaler.transform(feat_df[FEATURE_COLS])
    feat_df = feat_df.copy()
    feat_df["risk_probability"] = model.predict_proba(X_score_s)[:, 1]

    factor = score_params["pdo"] / np.log(2)
    offset = score_params["base_score"] - factor * np.log(score_params["base_odds"])
    def prob_to_score(p):
        p = np.clip(p, 1e-6, 1 - 1e-6)
        return np.clip(offset + factor * np.log((1 - p) / p), 300, 900)
    feat_df["credit_score"] = prob_to_score(feat_df["risk_probability"]).round(0)
    feat_df["risk_flag"] = pd.cut(feat_df["credit_score"], bins=[0, 580, 700, 900],
                                    labels=["RED", "AMBER", "GREEN"])

    contributions = X_score_s * model.coef_[0]
    contrib_df = pd.DataFrame(contributions, columns=FEATURE_COLS, index=feat_df.index)

    def top_reasons(idx, n=3):
        row = contrib_df.loc[idx]
        row_sorted = row.reindex(row.abs().sort_values(ascending=False).index)
        return [{"factor": REASON_LABELS[f], "direction": "increases_risk" if v > 0 else "reduces_risk", "weight": round(float(v), 3)}
                for f, v in row_sorted.head(n).items()]
    feat_df["top_reasons"] = [top_reasons(i) for i in feat_df.index]
    feat_df["scoring_method"] = "Statistical"
    return feat_df


def cold_start_score(dealers: pd.DataFrame, txns: pd.DataFrame, insufficient_ids: list) -> pd.DataFrame:
    """Rule-based provisional scoring for dealers with insufficient history."""
    if not insufficient_ids:
        return pd.DataFrame(columns=["dealer_id", "dealer_name", "city", "sector", "salesman_id", "salesman_name",
                                       "is_salesman_favorite", "credit_limit_pkr", "credit_score",
                                       "risk_flag", "risk_probability", "scoring_method", "top_reasons"])
    dealer_bounce = txns.groupby("dealer_id")["cheque_bounced"].mean()
    dwr = dealers.set_index("dealer_id").join(dealer_bounce.rename("bounce_rate"))
    dwr["bounce_rate"] = dwr["bounce_rate"].fillna(0)
    salesman_track = dwr.groupby("salesman_id")["bounce_rate"].mean()
    territory_base = dwr.groupby("territory_risk_tier")["bounce_rate"].mean()
    sector_base = dwr.groupby("sector")["bounce_rate"].mean()

    all_blended = []
    for _, d in dealers.iterrows():
        sr = salesman_track.get(d["salesman_id"], dealer_bounce.mean())
        tr = territory_base.get(d["territory_risk_tier"], dealer_bounce.mean())
        secr = sector_base.get(d["sector"], dealer_bounce.mean())
        all_blended.append(0.5 * sr + 0.3 * tr + 0.2 * secr)
    bmin, bmax = min(all_blended), max(all_blended)

    rows = []
    for dealer_id in insufficient_ids:
        d = dealers[dealers["dealer_id"] == dealer_id].iloc[0]
        sr = salesman_track.get(d["salesman_id"], dealer_bounce.mean())
        tr = territory_base.get(d["territory_risk_tier"], dealer_bounce.mean())
        secr = sector_base.get(d["sector"], dealer_bounce.mean())
        blended = 0.5 * sr + 0.3 * tr + 0.2 * secr
        norm = (blended - bmin) / (bmax - bmin + 1e-9)
        score = np.clip(680 - (norm * 200), 480, 680)
        tier = "AMBER-CAUTION" if score < 560 else "AMBER-STANDARD"
        rows.append({
            "dealer_id": dealer_id, "dealer_name": d["dealer_name"], "city": d["city"], "sector": d["sector"],
            "salesman_id": d["salesman_id"], "salesman_name": None, "is_salesman_favorite": bool(d["is_salesman_favorite"]),
            "credit_limit_pkr": d["credit_limit_pkr"], "credit_score": round(score), "risk_flag": tier,
            "risk_probability": None, "scoring_method": "Provisional (Cold-Start)",
            "top_reasons": [{"factor": "Insufficient payment history — score based on salesman/territory/sector averages", "direction": None, "weight": None}],
        })
    return pd.DataFrame(rows)

NAVY = RGBColor(0x1E, 0x27, 0x61)
RED = RGBColor(0xC0, 0x39, 0x2B)
AMBER = RGBColor(0xB8, 0x86, 0x0B)
GREEN = RGBColor(0x1E, 0x84, 0x49)
GREY = RGBColor(0x66, 0x66, 0x66)
DARK = RGBColor(0x1A, 0x1A, 0x1A)

TIER_COLOR = {"RED": RED, "AMBER": AMBER, "GREEN": GREEN}
TIER_LABEL = {"RED": "RED — High Risk", "AMBER": "AMBER — Moderate", "GREEN": "GREEN — Reliable"}
ACTION = {
    "RED": "Do not increase credit limit. Consider requiring partial upfront payment or cash-on-delivery terms on future orders.",
    "AMBER": "Maintain current credit limit. Monitor closely and reassess in 3 months.",
    "GREEN": "Eligible for credit limit review or increase based on strong payment history.",
}


def set_cell_background(cell, hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def fix_zoom_element(doc):
    """python-docx's default template ships a <w:zoom> element missing the
    required w:percent attribute -- present even in a totally blank document.
    Fixed here rather than left as a validation warning."""
    settings = doc.settings.element
    zoom = settings.find(qn("w:zoom"))
    if zoom is not None:
        zoom.set(qn("w:percent"), "100")

    # Same default-template issue: compatibilityMode is hardcoded to 14
    # (Word 2010), which triggers Word's "Compatibility Mode" banner on
    # every generated file. Bump to 15 (Word 2013+) to match modern Word.
    compat = settings.find(qn("w:compat"))
    if compat is not None:
        for cs in compat.findall(qn("w:compatSetting")):
            if cs.get(qn("w:name")) == "compatibilityMode":
                cs.set(qn("w:val"), "15")


def remove_table_borders(table):
    """Explicitly strips borders rather than relying on default table-style
    behavior, which can render visible gridlines depending on the Word/
    LibreOffice version -- guarantees the clean, borderless look regardless.
    OOXML's CT_TblPrBase schema requires child elements in a strict order;
    tblBorders must be inserted before shd/tblLayout/tblCellMar/tblLook if
    those already exist, not just appended at the end."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    following_tags = [qn(f"w:{t}") for t in
                       ("shd", "tblLayout", "tblCellMar", "tblLook", "tblCaption", "tblDescription")]
    insert_before = None
    for child in tbl_pr:
        if child.tag in following_tags:
            insert_before = child
            break
    if insert_before is not None:
        insert_before.addprevious(borders)
    else:
        tbl_pr.append(borders)

def tier_of(risk_flag: str) -> str:
    if risk_flag.startswith("RED"):
        return "RED"
    if risk_flag.startswith("GREEN"):
        return "GREEN"
    return "AMBER"
    
def build_risk_card_docx(dealer: dict) -> bytes:
    tier = tier_of(dealer["risk_flag"])
    color = TIER_COLOR[tier]
    is_contradiction = dealer.get("is_salesman_favorite") and tier == "RED"

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    letterhead = doc.add_paragraph()
    r = letterhead.add_run("[ DISTRIBUTOR LETTERHEAD ]")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY

    title = doc.add_paragraph()
    r = title.add_run("Dealer Credit Risk Card")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = DARK

    sub = doc.add_paragraph()
    r = sub.add_run("Generated from historical payment data — internal decision support only")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    doc.add_paragraph()

    id_table = doc.add_table(rows=0, cols=2)
    remove_table_borders(id_table)
    id_rows = [
        ("Dealer Name", dealer["dealer_name"]),
        ("City / Territory", dealer["city"]),
        ("Sector", dealer["sector"]),
        ("Assigned Salesman", f"{dealer.get('salesman_name') or dealer['salesman_id']} ({dealer['salesman_id']})"),
    ]
    if is_contradiction:
        id_rows.append(("Salesman Relationship", "Marked as a trusted / favorite account"))
    for label, value in id_rows:
        row = id_table.add_row()
        row.cells[0].width = Inches(2.2)
        row.cells[0].paragraphs[0].add_run(label).font.size = Pt(10)
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = GREY
        run = row.cells[1].paragraphs[0].add_run(str(value))
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RED if label == "Salesman Relationship" else DARK

    doc.add_paragraph()

    score_table = doc.add_table(rows=1, cols=2)
    score_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    remove_table_borders(score_table)
    left_cell, right_cell = score_table.rows[0].cells
    set_cell_background(left_cell, "F2F2F2")
    set_cell_background(right_cell, "F2F2F2")

    p = left_cell.paragraphs[0]
    r = p.add_run("CREDIT SCORE\n")
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    r2 = p.add_run(f"{round(dealer['credit_score'])}\n")
    r2.bold = True
    r2.font.size = Pt(34)
    r2.font.color.rgb = color
    r3 = p.add_run("out of 300–900")
    r3.font.size = Pt(8)
    r3.font.color.rgb = GREY

    p2 = right_cell.paragraphs[0]
    r4 = p2.add_run("RISK FLAG\n")
    r4.font.size = Pt(9)
    r4.font.color.rgb = GREY
    r5 = p2.add_run(f"{TIER_LABEL[tier]}\n")
    r5.bold = True
    r5.font.size = Pt(16)
    r5.font.color.rgb = color
    conf = dealer.get("risk_probability")
    conf_text = (f"Model confidence (ranking, not exact probability): {conf*100:.1f}%"
                 if conf is not None else
                 "Provisional score — based on portfolio averages, not this dealer's own history")
    r6 = p2.add_run(conf_text)
    r6.font.size = Pt(8)
    r6.font.color.rgb = GREY

    doc.add_paragraph()

    h = doc.add_paragraph()
    r = h.add_run("Key Contributing Factors")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = DARK

    for i, reason in enumerate(dealer.get("top_reasons", []), 1):
        direction = reason.get("direction")
        prefix = "Increases risk" if direction == "increases_risk" else "Reduces risk" if direction == "reduces_risk" else ""
        p = doc.add_paragraph()
        r = p.add_run(f"{i}. {prefix + ' — ' if prefix else ''}{reason['factor']}")
        r.bold = True
        r.font.size = Pt(10.5)

    doc.add_paragraph()
    h2 = doc.add_paragraph()
    r = h2.add_run("Recommended Action")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = DARK

    action_text = ACTION[tier]
    if is_contradiction:
        action_text += (" Recommend a direct conversation with the assigned salesman given this "
                         "account's trusted status contradicts the payment record.")
    p = doc.add_paragraph()
    p.add_run(action_text).font.size = Pt(10.5)

    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    r = disclaimer.add_run(
        "This score is generated from historical payment behavior only and does not replace "
        "human judgment. Risk tier (RED/AMBER/GREEN) reflects the model's relative ranking "
        "within this portfolio; the exact numeric score will sharpen in precision as more "
        "historical data becomes available."
    )
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = GREY

    buf = io.BytesIO()
    fix_zoom_element(doc)
    doc.save(buf)
    return buf.getvalue()

# --- Input normalization ---------------------------------------------------
# The synthetic dataset carries columns a real distributor's export never
# would: territory_risk_tier is a risk grade this project assigned, and
# is_salesman_favorite is an annotation only the distributor can supply.
# Hard-requiring them meant real data was rejected outright. These helpers
# fill sensible substitutes and report what they did, so nothing is silent.

# --- Column aliasing -------------------------------------------------------
# The transactions file already accepted real-world headers, but dealers and
# salesmen demanded exact canonical names -- so an export saying "Dealer Code"
# was accepted in one file and rejected in another. That asymmetry was the
# single biggest barrier to a real distributor using this at all.

def _canon(col) -> str:
    """Collapses casing, spacing and punctuation so 'Credit Limit (Rs)' and
    'credit_limit_rs' compare equal."""
    s = re.sub(r"[\s\-\.]+", "_", str(col).strip().lower())
    s = re.sub(r"[^a-z0-9_]", "", s)
    return re.sub(r"_+", "_", s).strip("_")


DEALER_COLUMN_ALIASES = {
    "dealer_id": {"dealer_id", "dealer_code", "dealerid", "dealer_no", "customer_code",
                  "customer_id", "account_code", "account_no", "party_code", "party_id",
                  "code", "dealer"},
    "dealer_name": {"dealer_name", "customer_name", "party_name", "shop_name",
                    "account_name", "business_name", "firm_name", "name"},
    "city": {"city", "town", "location", "area", "station"},
    "sector": {"sector", "category", "industry", "business_type", "segment",
               "product_category"},
    "salesman_id": {"salesman_id", "salesman_code", "salesmanid", "salesman", "sales_rep",
                    "rep_code", "so_code", "sales_officer", "officer_code", "booker",
                    "booker_code"},
    "credit_limit_pkr": {"credit_limit_pkr", "credit_limit_rs", "credit_limit",
                         "creditlimit", "limit", "credit_limit_amount", "approved_limit"},
    "onboarding_date": {"onboarding_date", "date_added", "created", "created_on", "since",
                        "start_date", "registration_date", "account_opened",
                        "opening_date", "first_invoice_date"},
    "is_salesman_favorite": {"is_salesman_favorite", "favorite", "favourite", "trusted",
                             "preferred", "trusted_account", "key_account"},
    "territory_risk_tier": {"territory_risk_tier", "territory_risk", "risk_tier",
                            "zone_risk", "risk_grade"},
}

SALESMAN_COLUMN_ALIASES = {
    "salesman_id": {"salesman_id", "salesman_code", "salesmanid", "code", "so_code",
                    "rep_code", "officer_code", "booker_code", "booker", "officer",
                    "rep", "id"},
    "salesman_name": {"salesman_name", "name", "salesman", "rep_name", "officer_name",
                      "booker_name", "full_name", "employee_name"},
}

TRANSACTION_COLUMN_ALIASES = {
    "transaction_id": {"transaction_id", "txn_id", "txnid", "invoice_no", "invoice_number",
                       "invoice_id", "bill_no", "bill_number", "voucher_no", "doc_no",
                       "document_no", "ref_no", "reference_no"},
    "dealer_id": {"dealer_id", "dealer_code", "dealerid", "customer_code", "customer_id",
                  "party_code", "party_id", "account_code", "account_no", "code", "dealer",
                  "customer", "party"},
    "invoice_date": {"invoice_date", "invoice_dt", "bill_date", "doc_date", "document_date",
                     "voucher_date", "issue_date", "date"},
    "due_date": {"due_date", "due_dt", "due", "maturity_date", "payment_due",
                 "payment_due_date"},
    "payment_date": {"payment_date", "paid_on", "paid_date", "receipt_date", "settled_on",
                     "clearing_date", "cleared_on", "recovery_date"},
    "amount_pkr": {"amount_pkr", "amount_rs", "amount", "net_amount", "invoice_amount",
                   "bill_amount", "gross_amount", "value", "total", "total_amount"},
    "payment_method": {"payment_method", "payment_mode", "mode", "mode_of_payment",
                       "pay_mode", "instrument", "instrument_type"},
    "cheque_bounced": {"cheque_bounced", "cheque_returned", "bounced", "bounce", "returned",
                       "dishonoured", "dishonored", "cheque_bounce", "is_bounced"},
}

def apply_column_aliases(df: pd.DataFrame, alias_map: dict, label: str):
    """Renames recognised header variants to canonical names. A column that
    already carries the canonical name is never clobbered by an alias."""
    canon_to_actual = {}
    for actual in df.columns:
        canon_to_actual.setdefault(_canon(actual), actual)

    rename, notes = {}, []
    for target, variants in alias_map.items():
        if target in df.columns:
            continue
        for v in variants:
            actual = canon_to_actual.get(v)
            if actual is None or actual in rename:
                continue
            rename[actual] = target
            if _canon(actual) != target:
                notes.append(f"{label}: read column '{actual}' as '{target}'")
            break
    return df.rename(columns=rename), notes

ESSENTIAL_DEALER_COLUMNS = ["dealer_id", "salesman_id", "credit_limit_pkr", "onboarding_date"]


def normalize_dealers(dealers: pd.DataFrame) -> tuple[pd.DataFrame, list[str]]:
    """Returns (normalized_dealers, notes). Raises ValueError if a genuinely
    un-substitutable column is missing."""
    df, notes = apply_column_aliases(dealers.copy(), DEALER_COLUMN_ALIASES, "dealers")

    missing = [c for c in ESSENTIAL_DEALER_COLUMNS if c not in df.columns]
    if missing:
        raise ValueError(
            "dealers file is missing required column(s): " + ", ".join(missing)
        )

    # Credit limits arrive as strings in real exports -- "Rs. 150,000",
    # "250,000/-", padded whitespace. The transactions pipeline already parses
    # these; the dealers file did not, so a real export crashed with a
    # TypeError deep inside feature computation.
    df["credit_limit_pkr"] = df["credit_limit_pkr"].apply(parse_messy_amount)
    if df["credit_limit_pkr"].isna().any():
        bad = df.loc[df["credit_limit_pkr"].isna(), "dealer_id"].head(5).tolist()
        raise ValueError(
            "Could not read the credit limit for some dealers "
            f"(e.g. {', '.join(map(str, bad))}). Accepted formats include "
            "150000, 'Rs. 150,000' and '150,000/-'."
        )

    if "dealer_name" not in df.columns:
        df["dealer_name"] = df["dealer_id"].astype(str)
        notes.append("dealer_name not provided — using dealer_id")
    if "city" not in df.columns:
        df["city"] = "Unknown"
        notes.append("city not provided — defaulted to 'Unknown'")
    if "sector" not in df.columns:
        df["sector"] = "Unspecified"
        notes.append("sector not provided — defaulted to 'Unspecified'")

    if "is_salesman_favorite" not in df.columns:
        df["is_salesman_favorite"] = False
        notes.append(
            "is_salesman_favorite not provided — defaulted to False. The "
            "trusted-but-risky contradiction flag needs this judgement from "
            "the distributor and will not fire."
        )
    else:
        # A real export may use Yes/No or 1/0 rather than True/False.
        df["is_salesman_favorite"] = (
            df["is_salesman_favorite"].astype(str).str.strip().str.lower()
            .isin({"true", "yes", "y", "1"})
        )

    if "territory_risk_tier" not in df.columns:
        df["territory_risk_tier"] = df["city"]
        notes.append(
            "territory_risk_tier not provided — grouping territory risk by "
            "city instead (real granularity rather than an invented grade)"
        )

    return df, notes


def normalize_salesmen(salesmen: pd.DataFrame) -> tuple[pd.DataFrame, list[str]]:
    """salesman_name is display-only; only salesman_id is truly required."""
    df, notes = apply_column_aliases(salesmen.copy(), SALESMAN_COLUMN_ALIASES, "salesmen")

    if "salesman_id" not in df.columns:
        raise ValueError("salesmen file is missing required column: salesman_id")
    if "salesman_name" not in df.columns:
        df["salesman_name"] = df["salesman_id"].astype(str)
        notes.append("salesman_name not provided — using salesman_id")

    return df, notes

# --- Model applicability diagnostics ---------------------------------------
# The model's scaler was fit on one distributor's portfolio. On a very
# different portfolio the scores still RANK correctly but can compress toward
# the 300/900 bounds, losing discrimination -- silently. These checks make
# that visible. Both are READ-ONLY: they observe scores, never alter them.
#
# Thresholds are heuristic, chosen so a portfolio resembling the training data
# reads "low" and a clearly divergent one reads "high". Tune if needed.

SHIFT_MODERATE_SDS = 1.0
SHIFT_HIGH_SDS = 2.5
SATURATION_WARN_PCT = 0.10


def assess_distribution_shift(feat_df: pd.DataFrame, artifact: dict) -> dict:
    """How far this portfolio's feature means sit from the model's training
    distribution, measured in training standard deviations."""
    scaler = artifact["scaler"]
    cols = artifact["feature_columns"]
    per_feature, max_shift = {}, 0.0

    for i, c in enumerate(cols):
        if c not in feat_df.columns:
            continue
        std = float(scaler.scale_[i]) or 1e-9
        disp = abs(float(feat_df[c].mean()) - float(scaler.mean_[i])) / std
        per_feature[c] = round(disp, 2)
        max_shift = max(max_shift, disp)

    if max_shift < SHIFT_MODERATE_SDS:
        severity = "low"
    elif max_shift < SHIFT_HIGH_SDS:
        severity = "moderate"
    else:
        severity = "high"

    return {
        "severity": severity,
        "max_displacement_sds": round(max_shift, 2),
        "per_feature_displacement_sds": per_feature,
    }


def assess_score_saturation(scored_df: pd.DataFrame) -> dict:
    """Fraction of dealers pinned at the 300/900 bounds. A high value means
    the model cannot tell those dealers apart -- the concrete symptom of an
    out-of-distribution portfolio."""
    n = len(scored_df)
    if n == 0:
        return {"severity": "low", "clipped_at_floor": 0,
                "clipped_at_ceiling": 0, "pct_clipped": 0.0}

    floor = int((scored_df["credit_score"] <= 300).sum())
    ceiling = int((scored_df["credit_score"] >= 900).sum())
    pct = (floor + ceiling) / n

    return {
        "severity": "high" if pct > SATURATION_WARN_PCT else "low",
        "clipped_at_floor": floor,
        "clipped_at_ceiling": ceiling,
        "pct_clipped": round(pct, 4),
    }


def build_reliability_report(feat_df: pd.DataFrame, scored_df: pd.DataFrame,
                             artifact: dict, model_source: str = "pretrained",
                             training_report: dict | None = None) -> dict:
    """Combines both checks into one verdict with plain-language guidance.

    model_source matters more than it looks. When a model has been trained on
    the uploaded portfolio, the distribution-shift check is measured against a
    scaler fitted to that same portfolio -- so it is ~0 by construction and
    tells you nothing. Reporting "this portfolio sits within the range the
    model was trained on" in that case is circular. What actually matters
    there is the new model's own cross-validated accuracy, plus whether the
    scores still saturate.
    """
    shift = assess_distribution_shift(feat_df, artifact)
    saturation = assess_score_saturation(scored_df)
    trained_here = model_source == "trained_on_your_data"

    if trained_here:
        auc = (training_report or {}).get("cv_auc_mean")
        auc_txt = (f" It scored AUC {auc:.3f} in cross-validation before being used."
                   if isinstance(auc, (int, float)) else "")
        if saturation["severity"] == "high":
            verdict = "use_ranking_only"
            guidance = (
                "Scored with a model trained on this portfolio's own payment "
                "history. A large share of dealers still sit at the extreme ends "
                "of the scale, so the numeric scores cannot separate them -- use "
                "the RED / AMBER / GREEN ranking instead." + auc_txt
            )
        else:
            verdict = "scores_reliable"
            guidance = (
                "Scored with a model trained on this portfolio's own payment "
                "history rather than the general model, because this portfolio "
                "differs from the data that model was built on." + auc_txt
            )
    elif shift["severity"] == "high" or saturation["severity"] == "high":
        verdict = "use_ranking_only"
        guidance = (
            "This portfolio differs substantially from the data the model was "
            "trained on. The RED / AMBER / GREEN ranking remains meaningful, but "
            "the numeric scores should not be read as absolute risk levels. For "
            "accurate absolute scoring, retrain the model on this distributor's "
            "own payment history."
        )
    elif shift["severity"] == "moderate":
        verdict = "scores_indicative"
        guidance = (
            "This portfolio differs somewhat from the model's training data. "
            "Rankings are reliable; treat exact score values as indicative."
        )
    else:
        verdict = "scores_reliable"
        guidance = (
            "This portfolio sits within the range the model was trained on. "
            "Scores and rankings can both be used as intended."
        )

    return {
        "verdict": verdict,
        "guidance": guidance,
        "distribution_shift": shift,
        "score_saturation": saturation,
        "shift_check_meaningful": not trained_here,
    }

# --- Runtime training on the uploaded portfolio ----------------------------
# A model trained on one distributor cannot be perfectly accurate for another.
# When the uploaded portfolio genuinely differs, the better answer is to train
# on THAT portfolio -- but only when the data actually supports it. Thin data,
# or a portfolio where nobody defaulted, will train "successfully" and produce
# confident nonsense, which is worse than an honest warning. Every gate below
# exists to prevent that, and the resulting model is cross-validated before
# it is ever allowed to score anything.

from sklearn.linear_model import LogisticRegression
from sklearn.preprocessing import StandardScaler
from sklearn.model_selection import StratifiedKFold
from sklearn.metrics import roc_auc_score

MIN_TRAINING_DEALERS = 50
MIN_MINORITY_CLASS = 10
MIN_SPAN_MONTHS = 18
MIN_ACCEPTABLE_CV_AUC = 0.65
# What counts as "paid late" when deriving training labels. A distributor on
# 60-day terms has entirely different norms than one on 15-day terms, so this
# cannot be a universal constant.
LATE_LABEL_THRESHOLD_DAYS = int(os.environ.get("LATE_PAYMENT_THRESHOLD_DAYS", 15))


def _resolve_training_windows(txns: pd.DataFrame):
    """Splits the uploaded history into a feature window and a later outcome
    window, so labels (who actually defaulted) can be derived."""
    d = txns["due_date"].dropna()
    if d.empty:
        return None, "no valid transaction dates"
    lo, hi = d.min(), d.max()
    span_months = (hi - lo).days / 30.44
    if span_months < MIN_SPAN_MONTHS:
        return None, (f"history spans only {span_months:.1f} months; at least "
                      f"{MIN_SPAN_MONTHS} are needed to separate a training window "
                      f"from an outcome window")
    label_months = max(6, min(12, span_months / 3))
    cutoff = hi - pd.DateOffset(months=int(round(label_months)))
    return {"cutoff": cutoff, "label_end": hi,
            "span_months": round(span_months, 1),
            "label_months": int(round(label_months))}, None

def resolve_late_threshold(label_window_txns: pd.DataFrame) -> tuple[float, dict]:
    """How many days late counts as a default -- derived per portfolio.

    This is a property of the market, not a universal constant. A distributor
    on 60-day terms has entirely different norms from one on 15-day terms; a
    single fixed number either labels the whole book as defaulters or none of
    it. In testing, a wholesale portfolio where every dealer paid 34-78 days
    late was labelled 100% high-risk by a 15-day rule, and training correctly
    refused to run on a degenerate target.

    Derived as median + 2 x MAD of this portfolio's own average lateness,
    floored at LATE_LABEL_THRESHOLD_DAYS. Median/MAD is robust to the very
    outliers being detected, and unlike a percentile it does not assume a
    fixed share of every book is bad: a tightly-clustered market flags few
    dealers, a widely-spread one flags more. A normal fast-paying market
    falls below the floor and keeps the conventional threshold.
    """
    floor = float(LATE_LABEL_THRESHOLD_DAYS)
    nb = label_window_txns[label_window_txns["cheque_bounced"] == False]
    if len(nb) == 0:
        return floor, {"days": floor,
                       "basis": "configured floor (no settled payments to measure)"}

    per_dealer = nb.groupby("dealer_id")["days_late"].mean().dropna()
    if len(per_dealer) < 5:
        return floor, {"days": floor,
                       "basis": "configured floor (too few dealers to infer a market norm)"}

    vals = per_dealer.values
    median = float(np.median(vals))
    mad = float(np.median(np.abs(vals - median)))
    derived = median + 2.0 * mad

    if derived <= floor:
        return floor, {"days": round(floor, 1),
                       "basis": f"configured floor (this market's norm is {median:.0f} days, "
                                f"below the {floor:.0f}-day floor)"}
    return derived, {"days": round(derived, 1),
                     "basis": f"derived from this portfolio (typical dealer pays {median:.0f} days "
                              f"late, so {derived:.0f} days marks an outlier)"}

def _build_training_labels(txns, cutoff, label_end):
    """Returns (labels_df, late_threshold_info)."""
    lw = txns[(txns["due_date"] >= cutoff) & (txns["due_date"] <= label_end)]
    threshold, late_info = resolve_late_threshold(lw)
    rows = []
    for did, g in lw.groupby("dealer_id"):
        if len(g) < MIN_INVOICES:
            continue
        nb = g[g["cheque_bounced"] == False]
        avg_late = nb["days_late"].mean() if len(nb) else 0
        rows.append({
            "dealer_id": did,
            "is_high_risk": int(bool(g["cheque_bounced"].any())
                                or avg_late > threshold),
        })
    return pd.DataFrame(rows), late_info


def _ks_statistic(y_true, y_prob):
    d = pd.DataFrame({"y": y_true, "p": y_prob}).sort_values("p")
    return float(np.max(np.abs(
        (d["y"] == 1).cumsum() / max((d["y"] == 1).sum(), 1)
        - (d["y"] == 0).cumsum() / max((d["y"] == 0).sum(), 1))))


def try_train_on_uploaded_data(dealers, salesmen, txns, reference_artifact):
    """Attempts a portfolio-specific model. Returns (artifact_or_None, report).
    None means fall back to the pretrained model -- the report says why."""
    feature_cols = reference_artifact["feature_columns"]
    score_params = reference_artifact["score_params"]

    windows, err = _resolve_training_windows(txns)
    if err:
        return None, {"attempted": True, "trained": False, "reason": err}

    feat_df, _ = compute_features(dealers, salesmen, txns,
                                  cutoff_date=windows["cutoff"])
    if len(feat_df) == 0:
        return None, {"attempted": True, "trained": False,
                      "reason": "no dealers had enough history before the training cutoff"}

    labels, late_info = _build_training_labels(txns, windows["cutoff"], windows["label_end"])
    if len(labels) == 0:
        return None, {"attempted": True, "trained": False,
                      "reason": "no dealers had enough activity in the outcome window"}

    pool = feat_df.merge(labels, on="dealer_id", how="inner")

    if len(pool) < MIN_TRAINING_DEALERS:
        return None, {"attempted": True, "trained": False,
                      "reason": (f"only {len(pool)} dealers have both training and "
                                 f"outcome history; at least {MIN_TRAINING_DEALERS} "
                                 f"are needed to train without overfitting")}

    counts = pool["is_high_risk"].value_counts()
    minority = int(counts.min()) if len(counts) == 2 else 0
    if minority < MIN_MINORITY_CLASS:
        return None, {"attempted": True, "trained": False,
                      "reason": (f"outcome classes too imbalanced "
                                 f"({ {int(k): int(v) for k, v in counts.items()} }); "
                                 f"need at least {MIN_MINORITY_CLASS} dealers in each "
                                 f"of defaulted / did-not-default")}

    X, y = pool[feature_cols], pool["is_high_risk"]
    aucs, kss = [], []
    for tr, te in StratifiedKFold(n_splits=5, shuffle=True,
                                  random_state=42).split(X, y):
        sc = StandardScaler().fit(X.iloc[tr])
        m = LogisticRegression(max_iter=1000, class_weight="balanced").fit(
            sc.transform(X.iloc[tr]), y.iloc[tr])
        p = m.predict_proba(sc.transform(X.iloc[te]))[:, 1]
        aucs.append(roc_auc_score(y.iloc[te], p))
        kss.append(_ks_statistic(y.iloc[te].values, p))

    auc_mean, auc_std = float(np.mean(aucs)), float(np.std(aucs))
    if auc_mean < MIN_ACCEPTABLE_CV_AUC:
        return None, {"attempted": True, "trained": False,
                      "reason": (f"a model trained on this portfolio scored only "
                                 f"AUC {auc_mean:.3f} in cross-validation, below the "
                                 f"{MIN_ACCEPTABLE_CV_AUC} minimum -- it would be "
                                 f"little better than guessing")}

    scaler = StandardScaler().fit(X)
    model = LogisticRegression(max_iter=1000, class_weight="balanced").fit(
        scaler.transform(X), y)

    metadata = {
        "source": "trained at runtime on uploaded portfolio",
        "training_pool_size": len(pool),
        "base_rate": round(float(y.mean()), 3),
        "cv_auc_mean": round(auc_mean, 3),
        "cv_auc_std": round(auc_std, 3),
        "cv_ks_mean": round(float(np.mean(kss)), 3),
        "feature_window_ends": str(windows["cutoff"].date()),
        "outcome_window_months": windows["label_months"],
        "history_span_months": windows["span_months"],
        "late_threshold_days": late_info["days"],
        "late_threshold_basis": late_info["basis"],
    }
    artifact = {"model": model, "scaler": scaler,
                "feature_columns": feature_cols, "score_params": score_params,
                "metadata": metadata}
    return artifact, {"attempted": True, "trained": True,
                      "reason": "portfolio had sufficient history and outcome signal",
                      **metadata}